import os
import io
import json
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Tuple

from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from google.oauth2.credentials import Credentials

from docx import Document
from docx.shared import Pt

__all__ = ["SimpleGoogleDrive"]

logger = logging.getLogger(__name__)

# -----------------------------
# Helpers
# -----------------------------
def _build_drive_service(credentials: Credentials):
    """Build Google Drive v3 service with discovery cache disabled (lower memory)."""
    return build("drive", "v3", credentials=credentials, cache_discovery=False)

def _safe_date(date_str: str) -> Optional[datetime]:
    try:
        return datetime.strptime(date_str, "%Y-%m-%d")
    except Exception:
        return None

def _escape_drive_name(value: str) -> str:
    """Make a name safe for a Drive v3 query single-quoted string (avoid plain apostrophes)."""
    return (value or "").replace("'", "’")

def _to_float(x) -> float:
    try:
        return float(x)
    except Exception:
        return 0.0

# -----------------------------
# Main class
# -----------------------------
class SimpleGoogleDrive:
    """
    Google Drive helper for WealthPro CRM.

    Root layouts supported:

    A) Letters directly under ROOT:
       ROOT
         A/
           Alice Smith/
         ...
         Z/

    B) Category folders under ROOT, then letters:
       ROOT
         Active Clients/
           A/
             Alice Smith/
         Archived Clients/
           A/
           ...

    Each client folder contains:
      - Tasks/
          Ongoing Tasks/
          Completed Tasks/
      - Reviews/
      - Products/                 (this file adds it)
    """

    def __init__(self, credentials: Credentials):
        self.drive = _build_drive_service(credentials)
        self.root_folder_id = os.environ.get("GDRIVE_ROOT_FOLDER_ID", "").strip()
        if not self.root_folder_id:
            raise RuntimeError("GDRIVE_ROOT_FOLDER_ID is not set. Please set it in Render env vars.")
        logger.info("Google Drive ready.")

    # -----------------------------
    # Low-level Drive ops
    # -----------------------------
    def _list_folders(self, parent_id: str) -> List[Dict]:
        """List non-trashed folders directly under parent."""
        folders: List[Dict] = []
        page_token = None
        query = (
            f"'{parent_id}' in parents and "
            "mimeType='application/vnd.google-apps.folder' and trashed=false"
        )
        while True:
            resp = self.drive.files().list(
                q=query,
                fields="nextPageToken, files(id, name)",
                pageToken=page_token,
                pageSize=1000,
            ).execute()
            folders.extend(resp.get("files", []))
            page_token = resp.get("nextPageToken")
            if not page_token:
                break
        return folders

    def _find_child_folder(self, parent_id: str, name: str) -> Optional[Dict]:
        """Find a folder named `name` directly under `parent_id`."""
        safe_name = _escape_drive_name(name)
        query = (
            f"'{parent_id}' in parents and "
            "mimeType='application/vnd.google-apps.folder' and "
            f"name='{safe_name}' and trashed=false"
        )
        resp = self.drive.files().list(q=query, fields="files(id, name)", pageSize=1).execute()
        files = resp.get("files", [])
        return files[0] if files else None

    def _ensure_folder(self, parent_id: str, name: str) -> str:
        """Get or create a child folder."""
        existing = self._find_child_folder(parent_id, name)
        if existing:
            return existing["id"]
        body = {
            "name": name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [parent_id],
        }
        created = self.drive.files().create(body=body, fields="id,name").execute()
        return created["id"]

    def _upload_bytes(self, parent_id: str, filename: str, data: bytes, mime: str) -> str:
        """Create a new file with bytes content."""
        media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime, resumable=False)
        body = {"name": filename, "parents": [parent_id]}
        created = self.drive.files().create(body=body, media_body=media, fields="id").execute()
        return created["id"]

    def _update_file_bytes(self, file_id: str, data: bytes, mime: str) -> None:
        """Replace the content of an existing file with bytes."""
        media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime, resumable=False)
        self.drive.files().update(fileId=file_id, media_body=media).execute()

    def _download_text(self, file_id: str) -> str:
        """Download a file as text (utf-8)."""
        request = self.drive.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        return fh.read().decode("utf-8", errors="ignore")

    def _move_file(self, file_id: str, new_parent_id: str):
        file = self.drive.files().get(fileId=file_id, fields="parents").execute()
        prev = ",".join(file.get("parents", [])) if file.get("parents") else ""
        self.drive.files().update(
            fileId=file_id, addParents=new_parent_id, removeParents=prev, fields="id,parents"
        ).execute()

    def _rename_file(self, file_id: str, new_name: str):
        self.drive.files().update(fileId=file_id, body={"name": new_name}, fields="id,name").execute()

    # -----------------------------
    # Folder discovery helpers
    # -----------------------------
    def _get_letter_folders(self, parent_id: str) -> List[Dict]:
        """Return A–Z (single uppercase letter) folders under parent."""
        out = []
        for f in self._list_folders(parent_id):
            nm = (f.get("name") or "").strip()
            if len(nm) == 1 and nm.isalpha() and nm.upper() == nm:
                out.append(f)
        return out

    def _has_client_markers(self, folder_id: str) -> bool:
        """Heuristic: treat a folder as a client if it contains 'Tasks' or 'Reviews' or 'Products'."""
        for f in self._list_folders(folder_id):
            nm = (f.get("name") or "").strip()
            if nm in {"Tasks", "Reviews", "Products"}:
                return True
        return False

    # -----------------------------
    # Client creation & listing
    # -----------------------------
    def create_client_enhanced_folders(self, display_name: str) -> str:
        """
        Create the client's A–Z index under the FIRST category that has letters,
        or directly under ROOT if letters are at root.
        Returns the client folder id.
        """
        display_name = (display_name or "").strip()
        if not display_name:
            raise ValueError("display_name required")

        # Prefer letters directly under ROOT
        root_letters = self._get_letter_folders(self.root_folder_id)

        parent_for_letters = None
        if root_letters:
            parent_for_letters = self.root_folder_id
        else:
            # Find a category (e.g., "Active Clients") that contains A–Z
            for cat in self._list_folders(self.root_folder_id):
                if self._get_letter_folders(cat["id"]):
                    parent_for_letters = cat["id"]
                    break
            if parent_for_letters is None:
                parent_for_letters = self.root_folder_id

        first = display_name[0].upper()
        index_letter = first if first.isalpha() else "#"
        index_id = self._ensure_folder(parent_for_letters, index_letter)

        client_id = self._ensure_folder(index_id, display_name)

        # Core structure
        tasks_id = self._ensure_folder(client_id, "Tasks")
        self._ensure_folder(tasks_id, "Ongoing Tasks")
        self._ensure_folder(tasks_id, "Completed Tasks")
        self._ensure_folder(client_id, "Reviews")
        self._ensure_folder(client_id, "Products")  # NEW

        logger.info("Created enhanced client folder for %s", display_name)
        return client_id

    def get_clients_enhanced(self) -> List[Dict]:
        """
        Discover client folders robustly and also populate each client's portfolio_value
        as the sum of their Products values (so dashboards & lists work without changing routes).
        """
        clients: List[Dict] = []

        def add_client(folder: Dict):
            cid = folder["id"]
            total_val = self._sum_products_for_client(cid)
            clients.append(
                {
                    "client_id": cid,
                    "display_name": (folder.get("name") or "").strip(),
                    "status": "active",
                    "folder_id": cid,
                    "portfolio_value": float(total_val),  # used by lists/dashboards
                }
            )

        # Case 1: letters directly under ROOT
        root_letters = self._get_letter_folders(self.root_folder_id)
        if root_letters:
            for letter in root_letters:
                for child in self._list_folders(letter["id"]):
                    add_client(child)
        else:
            # Case 2: categories under ROOT -> letters -> clients
            for category in self._list_folders(self.root_folder_id):
                letters = self._get_letter_folders(category["id"])
                if letters:
                    for letter in letters:
                        for child in self._list_folders(letter["id"]):
                            add_client(child)
                else:
                    if self._has_client_markers(category["id"]):
                        add_client(category)

        clients.sort(key=lambda c: (c["display_name"] or "").lower())
        return clients

    # -----------------------------
    # Tasks
    # -----------------------------
    def _get_client_tasks_folder_ids(self, client_id: str) -> Dict[str, str]:
        tasks_id = self._ensure_folder(client_id, "Tasks")
        ongoing_id = self._ensure_folder(tasks_id, "Ongoing Tasks")
        completed_id = self._ensure_folder(tasks_id, "Completed Tasks")
        return {"tasks": tasks_id, "ongoing": ongoing_id, "completed": completed_id}

    def add_task_enhanced(self, task: Dict, client: Dict) -> bool:
        """Save a task as a .txt file in Ongoing Tasks."""
        client_id = client.get("client_id") or client.get("folder_id")
        if not client_id:
            raise ValueError("client client_id/folder_id missing")

        fids = self._get_client_tasks_folder_ids(client_id)

        due = task.get("due_date", "")
        pr = task.get("priority", "Medium")
        ttype = task.get("task_type", "")
        title = (task.get("title") or "").strip()
        tid = task.get("task_id", f"TSK{datetime.now().strftime('%Y%m%d%H%M%S')}")

        filename = f"{due} - {pr} - {ttype} - {title} [{tid}].txt"

        lines = [
            f"Task ID: {tid}",
            f"Client ID: {client_id}",
            f"Title: {title}",
            f"Type: {ttype}",
            f"Priority: {pr}",
            f"Due Date: {due}",
            f"Status: {task.get('status','Pending')}",
            f"Created: {task.get('created_date','')}",
            f"Completed: {task.get('completed_date','')}",
        ]
        if task.get("time_spent"):
            lines.append(f"Time Allocated: {task.get('time_spent')}")
        if task.get("description"):
            lines.append("")
            lines.append("Description:")
            lines.append(task["description"])

        content = ("\n".join(lines)).encode("utf-8")
        self._upload_bytes(fids["ongoing"], filename, content, "text/plain")
        return True

    def complete_task(self, task_file_id: str) -> bool:
        """Move the task file to Completed Tasks and prefix with 'COMPLETED - '."""
        file = self.drive.files().get(fileId=task_file_id, fields="id,name,parents").execute()
        if not file:
            return False

        # climb up to find Tasks -> client
        parent = (file.get("parents") or [None])[0]
        client_id = None

        hops = 0
        while parent and hops < 5:
            node = self.drive.files().get(fileId=parent, fields="id,name,parents").execute()
            name = node.get("name") or ""
            if name == "Tasks":
                par = node.get("parents") or []
                client_id = par[0] if par else None
                break
            parent = (node.get("parents") or [None])[0]
            hops += 1

        if not client_id:
            return False

        fids = self._get_client_tasks_folder_ids(client_id)
        completed = fids["completed"]

        self._move_file(task_file_id, completed)

        current_name = file.get("name", "")
        if not current_name.startswith("COMPLETED - "):
            self._rename_file(task_file_id, f"COMPLETED - {current_name}")

        return True

    def _parse_task_filename(self, name: str) -> Dict:
        result = {"due_date": "", "priority": "", "task_type": "", "title": "", "task_id": ""}
        base = name[:-4] if name.lower().endswith(".txt") else name

        # Task ID in square brackets
        if "[" in base and "]" in base and base.rfind("[") < base.rfind("]"):
            s = base.rfind("[")
            e = base.rfind("]")
            result["task_id"] = base[s + 1 : e].strip()
            base = (base[:s] + base[e + 1 :]).strip()

        parts = [p.strip() for p in base.split(" - ", 3)]
        if len(parts) >= 4:
            result["due_date"], result["priority"], result["task_type"], result["title"] = parts
        else:
            if parts:
                result["due_date"] = parts[0]
            if len(parts) > 1:
                result["priority"] = parts[1]
            if len(parts) > 2:
                result["task_type"] = parts[2]
            if len(parts) > 3:
                result["title"] = parts[3]

        return result

    def get_client_tasks(self, client_id: str) -> List[Dict]:
        fids = self._get_client_tasks_folder_ids(client_id)
        out: List[Dict] = []

        for status, folder in (("Pending", fids["ongoing"]), ("Completed", fids["completed"])):  # type: ignore
            page = None
            while True:
                resp = self.drive.files().list(
                    q=(
                        f"'{folder}' in parents and "
                        "mimeType!='application/vnd.google-apps.folder' and trashed=false"
                    ),
                    fields="nextPageToken, files(id,name,createdTime,modifiedTime)",
                    pageToken=page,
                    orderBy="name_natural",
                ).execute()
                for f in resp.get("files", []):
                    meta = self._parse_task_filename(f.get("name", ""))
                    out.append(
                        {
                            "task_id": f.get("id"),
                            "client_id": client_id,
                            "title": meta.get("title", ""),
                            "task_type": meta.get("task_type", ""),
                            "due_date": meta.get("due_date", ""),
                            "priority": meta.get("priority", "Medium"),
                            "status": status,
                            "description": "",
                            "created_date": (f.get("createdTime", "")[:10] or ""),
                            "completed_date": (f.get("modifiedTime", "")[:10] if status == "Completed" else ""),
                            "time_spent": "",
                        }
                    )
                page = resp.get("nextPageToken")
                if not page:
                    break

        def sort_key(t):
            d = _safe_date(t.get("due_date", ""))
            return (0 if t["status"] == "Pending" else 1, d or datetime(1970, 1, 1))

        out.sort(key=sort_key)
        return out

    def get_upcoming_tasks(self, days: int = 30) -> List[Dict]:
        """Scan all clients' Ongoing Tasks and return those due within `days`."""
        upcoming: List[Dict] = []
        clients = self.get_clients_enhanced()
        today = datetime.today().date()
        horizon = today + timedelta(days=days)

        for c in clients:
            client_id = c["client_id"]
            fids = self._get_client_tasks_folder_ids(client_id)
            ongoing = fids["ongoing"]
            page = None
            while True:
                resp = self.drive.files().list(
                    q=(
                        f"'{ongoing}' in parents and "
                        "mimeType!='application/vnd.google-apps.folder' and trashed=false"
                    ),
                    fields="nextPageToken, files(id,name,createdTime)",
                    pageToken=page,
                    orderBy="name_natural",
                ).execute()
                for f in resp.get("files", []):
                    meta = self._parse_task_filename(f.get("name", ""))
                    dd = _safe_date(meta.get("due_date", ""))
                    if dd and today <= dd.date() <= horizon:
                        upcoming.append(
                            {
                                "task_id": f.get("id"),
                                "client_id": client_id,
                                "title": meta.get("title", ""),
                                "task_type": meta.get("task_type", ""),
                                "due_date": meta.get("due_date", ""),
                                "priority": meta.get("priority", "Medium"),
                                "status": "Pending",
                                "description": "",
                                "created_date": f.get("createdTime", "")[:10],
                                "completed_date": "",
                                "time_spent": "",
                            }
                        )
                page = resp.get("nextPageToken")
                if not page:
                    break

        upcoming.sort(key=lambda t: _safe_date(t["due_date"]) or datetime(1970, 1, 1))
        return upcoming

    # -----------------------------
    # Reviews (unchanged behavior)
    # -----------------------------
    def _uk_date_str(self, dt: datetime) -> str:
        try:
            return dt.strftime("%-d %B %Y")
        except ValueError:
            return dt.strftime("%d %B %Y")

    def create_review_pack_for_client(self, client: Dict) -> Dict[str, str]:
        """Build Review <YEAR> structure and create two Word docs in 'Agenda & Valuation'."""
        client_id = client.get("client_id") or client.get("folder_id")
        display_name = client.get("display_name") or "Client"
        if not client_id:
            raise ValueError("client client_id/folder_id missing")

        year = datetime.today().year
        reviews_root = self._ensure_folder(client_id, "Reviews")
        yr_id = self._ensure_folder(reviews_root, f"Review {year}")

        subfolders = [
            "Agenda & Valuation",
            "FF&ATR",
            "ID&V & Sanction Search",
            "Meeting Notes",
            "Research",
            "Review Letter",
            "Client Confirmation",
            "Emails",
        ]
        created = {"review_year_id": yr_id}
        for sf in subfolders:
            created[sf] = self._ensure_folder(yr_id, sf)

        agenda_val = created["Agenda & Valuation"]
        today_str = self._uk_date_str(datetime.today())

        # Agenda doc (basic styled)
        agenda_doc = self._build_agenda_doc(display_name, today_str)
        self._upload_docx(
            agenda_val,
            f"Meeting Agenda – {display_name} – {year}.docx",
            agenda_doc,
        )

        # Valuation doc (basic styled)
        val_doc = self._build_valuation_doc(display_name, today_str)
        self._upload_docx(
            agenda_val,
            f"Valuation Summary – {display_name} – {year}.docx",
            val_doc,
        )

        return created

    def _upload_docx(self, parent_id: str, filename: str, document: Document):
        stream = io.BytesIO()
        document.save(stream)
        stream.seek(0)
        media = MediaIoBaseUpload(
            stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=False,
        )
        meta = {"name": filename, "parents": [parent_id]}
        self.drive.files().create(body=meta, media_body=media, fields="id,name").execute()

    # Word doc builders (kept simple & consistent)
    def _build_agenda_doc(self, client_display_name: str, date_str: str) -> Document:
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Client Review Meeting Agenda")
        r.bold = True
        r.font.size = Pt(16)
        doc.add_paragraph(f"Client: {client_display_name}")
        doc.add_paragraph(f"Date: {date_str}")
        doc.add_paragraph("")
        items = [
            "1. Welcome & objectives",
            "2. Personal & financial updates",
            "3. Investment performance & valuation",
            "4. Risk profile (ATR) & capacity",
            "5. Charges & costs",
            "6. Portfolio changes / recommendations",
            "7. Action points & next steps",
        ]
        for it in items:
            doc.add_paragraph(it)
        return doc

    def _build_valuation_doc(self, client_display_name: str, date_str: str) -> Document:
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Valuation Summary")
        r.bold = True
        r.font.size = Pt(16)
        doc.add_paragraph(f"Client: {client_display_name}")
        doc.add_paragraph(f"Date: {date_str}")
        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Plan / Account"
        hdr[1].text = "Provider"
        hdr[2].text = "Value (£)"
        row = table.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        doc.add_paragraph("")
        doc.add_paragraph("Total Value: £")
        return doc

    # -----------------------------
    # Products (NEW)
    # -----------------------------
    def _get_client_products_folder(self, client_id: str) -> str:
        return self._ensure_folder(client_id, "Products")

    def add_product(self, product: Dict, client: Dict) -> bool:
        """
        Save a product JSON file in Products/.
        Fields expected:
        - company (str)
        - portfolio (str)
        - product_name (str)
        - value (float)
        - charge_pct (float)  (e.g., 1.0 means 1%)
        - as_of (YYYY-MM-DD)
        """
        client_id = client.get("client_id") or client.get("folder_id")
        if not client_id:
            raise ValueError("client client_id/folder_id missing")

        folder_id = self._get_client_products_folder(client_id)

        pid = product.get("product_id", f"PRD{datetime.now().strftime('%Y%m%d%H%M%S')}")
        company = (product.get("company") or "").strip()
        portfolio = (product.get("portfolio") or "").strip()
        name = (product.get("product_name") or "").strip()
        value = _to_float(product.get("value"))
        charge_pct = _to_float(product.get("charge_pct"))
        as_of = (product.get("as_of") or datetime.now().strftime("%Y-%m-%d")).strip()

        # Filename: YYYY-MM-DD - Company - Portfolio - Product [PRDxxxxxxxx].json
        filename = f"{as_of} - {company} - {portfolio} - {name} [{pid}].json"

        payload = {
            "product_id": pid,
            "company": company,
            "portfolio": portfolio,
            "product_name": name,
            "value": value,
            "charge_pct": charge_pct,
            "as_of": as_of,
            "client_id": client_id,
        }
        data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self._upload_bytes(folder_id, filename, data, "application/json")
        return True

    def update_product(self, file_id: str, product: Dict) -> bool:
        """Replace JSON content of an existing product file."""
        payload = {
            "product_id": product.get("product_id", ""),
            "company": (product.get("company") or "").strip(),
            "portfolio": (product.get("portfolio") or "").strip(),
            "product_name": (product.get("product_name") or "").strip(),
            "value": _to_float(product.get("value")),
            "charge_pct": _to_float(product.get("charge_pct")),
            "as_of": (product.get("as_of") or datetime.now().strftime("%Y-%m-%d")).strip(),
            "client_id": product.get("client_id", ""),
        }
        data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self._update_file_bytes(file_id, data, "application/json")
        return True

    def _parse_product_file(self, file_id: str, name: str) -> Optional[Dict]:
        """Read JSON product file and return dict with computed annual_fee."""
        try:
            txt = self._download_text(file_id)
            obj = json.loads(txt)
        except Exception as e:
            logger.error(f"Failed to parse product file {file_id}: {e}")
            return None

        value = _to_float(obj.get("value"))
        pct = _to_float(obj.get("charge_pct"))
        annual_fee = value * (pct / 100.0)

        return {
            "file_id": file_id,
            "product_id": obj.get("product_id", ""),
            "company": obj.get("company", ""),
            "portfolio": obj.get("portfolio", ""),
            "product_name": obj.get("product_name", ""),
            "value": value,
            "charge_pct": pct,
            "annual_fee": annual_fee,
            "as_of": obj.get("as_of", ""),
            "client_id": obj.get("client_id", ""),
            "file_name": name,
        }

    def get_client_products(self, client_id: str) -> List[Dict]:
        folder_id = self._get_client_products_folder(client_id)
        out: List[Dict] = []
        page = None
        while True:
            resp = self.drive.files().list(
                q=(
                    f"'{folder_id}' in parents and "
                    "mimeType!='application/vnd.google-apps.folder' and trashed=false"
                ),
                fields="nextPageToken, files(id,name,createdTime,modifiedTime,mimeType)",
                pageToken=page,
                orderBy="name_natural desc",
            ).execute()
            for f in resp.get("files", []):
                if not f.get("name", "").lower().endswith(".json"):
                    continue
                parsed = self._parse_product_file(f["id"], f.get("name", ""))
                if parsed:
                    out.append(parsed)
            page = resp.get("nextPageToken")
            if not page:
                break
        # newest first by as_of (encoded at start of filename) already handled by orderBy
        return out

    def _sum_products_for_client(self, client_id: str) -> float:
        prods = self.get_client_products(client_id)
        return float(sum(p.get("value", 0.0) for p in prods))

    def get_global_product_options(self) -> Tuple[List[str], List[str]]:
        """Return (companies, portfolios) unique lists across all clients."""
        companies = set()
        portfolios = set()

        # Iterate all clients, then their Products
        clients = []
        # Find all client folders similar to get_clients_enhanced but lighter
        root_letters = self._get_letter_folders(self.root_folder_id)
        if root_letters:
            for letter in root_letters:
                for child in self._list_folders(letter["id"]):
                    clients.append(child["id"])
        else:
            for category in self._list_folders(self.root_folder_id):
                letters = self._get_letter_folders(category["id"])
                if letters:
                    for letter in letters:
                        for child in self._list_folders(letter["id"]):
                            clients.append(child["id"])
                else:
                    if self._has_client_markers(category["id"]):
                        clients.append(category["id"])

        for cid in clients:
            for p in self.get_client_products(cid):
                c = (p.get("company") or "").strip()
                pf = (p.get("portfolio") or "").strip()
                if c:
                    companies.add(c)
                if pf:
                    portfolios.add(pf)

        return (sorted(companies), sorted(portfolios))

    def get_total_assets_from_products(self) -> float:
        """Sum of all product values across all clients."""
        total = 0.0
        # reuse get_clients_enhanced which now sets portfolio_value per client
        for c in self.get_clients_enhanced():
            total += float(c.get("portfolio_value") or 0.0)
        return float(total)
